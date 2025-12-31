"""
Knowledge Base service for managing prompts, documents, FAQs and RAG integration
"""
from typing import List, Optional, Tuple, Dict, Any
import logging
from sqlalchemy.orm import Session
from sqlalchemy import func, or_
from fastapi import UploadFile

from app.db.models.knowledge_base import (
    KBCategory, KBPrompt, KBDocument, KBDocumentChunk, KBFaq
)
from app.db.models.file import File
from app.services.ai_service import ai_service
from app.services.file_service import file_service
from app.core.config import settings

logger = logging.getLogger(__name__)


class KnowledgeBaseService:
    """Service for Knowledge Base operations"""

    # ============== Category Operations ==============

    async def create_category(self, db: Session, data: dict) -> KBCategory:
        """Create a new KB category"""
        category = KBCategory(**data)
        db.add(category)
        db.commit()
        db.refresh(category)
        return category

    async def get_categories(
        self, db: Session, include_inactive: bool = False,
        lounge_id: Optional[int] = None, include_global: bool = True
    ) -> List[KBCategory]:
        """Get all categories with counts, filtered by lounge"""
        query = db.query(KBCategory)
        if not include_inactive:
            query = query.filter(KBCategory.is_active == True)

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBCategory.lounge_id == lounge_id, KBCategory.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBCategory.lounge_id == lounge_id)
        elif not include_global:
            # Only global categories
            query = query.filter(KBCategory.lounge_id.is_(None))

        return query.order_by(KBCategory.sort_order, KBCategory.name).all()

    async def get_category(self, db: Session, category_id: int) -> Optional[KBCategory]:
        """Get a category by ID"""
        return db.query(KBCategory).filter(KBCategory.id == category_id).first()

    async def update_category(
        self, db: Session, category_id: int, data: dict
    ) -> Optional[KBCategory]:
        """Update a category"""
        category = await self.get_category(db, category_id)
        if not category:
            return None

        for key, value in data.items():
            if value is not None:
                setattr(category, key, value)

        db.commit()
        db.refresh(category)
        return category

    async def delete_category(self, db: Session, category_id: int) -> bool:
        """Delete a category"""
        category = await self.get_category(db, category_id)
        if not category:
            return False

        db.delete(category)
        db.commit()
        return True

    def get_category_counts(self, db: Session, category: KBCategory) -> dict:
        """Get content counts for a category"""
        prompt_count = db.query(func.count(KBPrompt.id)).filter(
            KBPrompt.category_id == category.id
        ).scalar()
        document_count = db.query(func.count(KBDocument.id)).filter(
            KBDocument.category_id == category.id
        ).scalar()
        faq_count = db.query(func.count(KBFaq.id)).filter(
            KBFaq.category_id == category.id
        ).scalar()

        return {
            "prompt_count": prompt_count or 0,
            "document_count": document_count or 0,
            "faq_count": faq_count or 0
        }

    # ============== Prompt Operations ==============

    async def create_prompt(
        self, db: Session, data: dict, admin_id: int, skip_embedding: bool = False
    ) -> KBPrompt:
        """Create a new prompt and optionally generate embedding"""
        prompt = KBPrompt(**data, created_by_id=admin_id)
        db.add(prompt)
        db.commit()
        db.refresh(prompt)

        # Generate embedding if included in RAG (unless skipped for background processing)
        if prompt.is_included_in_rag and not skip_embedding:
            await self._update_prompt_embedding(db, prompt)

        return prompt

    async def _update_prompt_embedding(
        self, db: Session, prompt: KBPrompt
    ) -> None:
        """Generate and store embedding for prompt"""
        try:
            text = f"{prompt.title}\n{prompt.content}"
            if prompt.description:
                text += f"\n{prompt.description}"

            embedding = await ai_service.create_embedding(text)
            prompt.embedding = embedding
            prompt.embedding_model = settings.OPENAI_EMBEDDING_MODEL
            db.commit()
        except Exception as e:
            logger.error(f"Error generating prompt embedding: {e}")

    async def get_prompts_paginated(
        self,
        db: Session,
        page: int = 1,
        limit: int = 20,
        category_id: Optional[int] = None,
        search: Optional[str] = None,
        is_active: Optional[bool] = None,
        lounge_id: Optional[int] = None,
        include_global: bool = True
    ) -> Tuple[List[KBPrompt], int]:
        """Get paginated prompts with filters"""
        query = db.query(KBPrompt)

        if category_id:
            query = query.filter(KBPrompt.category_id == category_id)
        if is_active is not None:
            query = query.filter(KBPrompt.is_active == is_active)
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                or_(
                    KBPrompt.title.ilike(search_term),
                    KBPrompt.content.ilike(search_term),
                    KBPrompt.description.ilike(search_term)
                )
            )

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBPrompt.lounge_id == lounge_id, KBPrompt.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBPrompt.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBPrompt.lounge_id.is_(None))

        total = query.count()
        skip = (page - 1) * limit
        prompts = query.order_by(KBPrompt.created_at.desc()).offset(skip).limit(limit).all()

        return prompts, total

    async def get_prompt(self, db: Session, prompt_id: int) -> Optional[KBPrompt]:
        """Get a prompt by ID"""
        return db.query(KBPrompt).filter(KBPrompt.id == prompt_id).first()

    async def update_prompt(
        self, db: Session, prompt_id: int, data: dict
    ) -> Optional[KBPrompt]:
        """Update a prompt"""
        prompt = await self.get_prompt(db, prompt_id)
        if not prompt:
            return None

        content_changed = False
        for key, value in data.items():
            if value is not None:
                if key in ['title', 'content', 'description']:
                    content_changed = True
                setattr(prompt, key, value)

        db.commit()
        db.refresh(prompt)

        # Regenerate embedding if content changed and RAG is enabled
        if content_changed and prompt.is_included_in_rag:
            await self._update_prompt_embedding(db, prompt)

        return prompt

    async def delete_prompt(self, db: Session, prompt_id: int) -> bool:
        """Delete a prompt"""
        prompt = await self.get_prompt(db, prompt_id)
        if not prompt:
            return False

        db.delete(prompt)
        db.commit()
        return True

    async def regenerate_prompt_embedding(
        self, db: Session, prompt_id: int
    ) -> Optional[KBPrompt]:
        """Regenerate embedding for a prompt"""
        prompt = await self.get_prompt(db, prompt_id)
        if not prompt:
            return None

        await self._update_prompt_embedding(db, prompt)
        db.refresh(prompt)
        return prompt

    # ============== Document Operations ==============

    async def create_document(
        self,
        db: Session,
        file: UploadFile,
        data: dict,
        admin_id: int
    ) -> KBDocument:
        """Upload document and create KB entry"""
        # Upload file using existing file_service
        file_record = await file_service.upload_file(
            file=file,
            user_id=admin_id,
            db=db,
            folder="knowledge_base/documents"
        )

        # Get file extension
        file_ext = file.filename.split('.')[-1].lower() if file.filename else 'unknown'

        # Create document record
        document = KBDocument(
            title=data.get('title', file.filename),
            description=data.get('description'),
            category_id=data.get('category_id'),
            tags=data.get('tags'),
            file_id=file_record.id,
            original_filename=file.filename,
            file_type=file_ext,
            file_size_bytes=file_record.size_bytes,
            created_by_id=admin_id
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # Process document in background (text extraction, chunking, embeddings)
        # For now, we'll do it synchronously - in production use Celery
        await self._process_document(db, document)

        return document

    async def _process_document(self, db: Session, document: KBDocument) -> None:
        """Process document: extract text, generate summary, create chunks"""
        try:
            # Text extraction based on file type
            extracted_text = await self._extract_text_from_document(db, document)

            if not extracted_text:
                document.processing_error = "Could not extract text from document"
                db.commit()
                return

            document.extracted_text = extracted_text

            # Generate summary
            if len(extracted_text) > 500:
                try:
                    summary = await ai_service.summarize_text(extracted_text[:10000])
                    document.summary = summary
                except Exception as e:
                    logger.error(f"Error generating summary: {e}")

            # Create chunks for RAG
            chunks = self._chunk_text(extracted_text)
            for i, chunk_text in enumerate(chunks):
                chunk = KBDocumentChunk(
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=i,
                    token_count=len(chunk_text.split())
                )
                db.add(chunk)

            db.commit()

            # Generate embeddings for chunks
            await self._generate_chunk_embeddings(db, document)

            # Generate document-level embedding
            await self._update_document_embedding(db, document)

            document.is_processed = True
            document.processing_error = None
            db.commit()

        except Exception as e:
            document.processing_error = str(e)
            db.commit()
            logger.error(f"Error processing document {document.id}: {e}")

    async def _extract_text_from_document(
        self, db: Session, document: KBDocument
    ) -> Optional[str]:
        """Extract text from document based on file type"""
        try:
            # Get file content from storage (works with both S3 and local)
            content = await file_service.get_file_content(document.file_id, db)
            if not content:
                logger.error(f"Could not retrieve file content for document {document.id}")
                return None

            if document.file_type == 'txt':
                # For TXT files, decode directly
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    return content.decode('latin-1')

            elif document.file_type == 'pdf':
                # Use PyPDF2 for PDF extraction
                try:
                    import io
                    from PyPDF2 import PdfReader

                    pdf_file = io.BytesIO(content)
                    reader = PdfReader(pdf_file)
                    text_parts = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return "\n\n".join(text_parts)
                except ImportError:
                    logger.warning("PyPDF2 not installed. Cannot extract PDF text.")
                    return None
                except Exception as e:
                    logger.error(f"Error extracting PDF text: {e}")
                    return None

            elif document.file_type in ['docx', 'doc']:
                # Use python-docx for DOCX extraction
                try:
                    import io
                    from docx import Document

                    docx_file = io.BytesIO(content)
                    doc = Document(docx_file)
                    text_parts = [para.text for para in doc.paragraphs if para.text]
                    return "\n\n".join(text_parts)
                except ImportError:
                    logger.warning("python-docx not installed. Cannot extract DOCX text.")
                    return None
                except Exception as e:
                    logger.error(f"Error extracting DOCX text: {e}")
                    return None

            return None
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return None

    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence ending
                for punct in ['. ', '! ', '? ', '\n\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct > start:
                        end = last_punct + len(punct)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks

    async def _generate_chunk_embeddings(
        self, db: Session, document: KBDocument
    ) -> None:
        """Generate embeddings for all chunks of a document"""
        chunks = db.query(KBDocumentChunk).filter(
            KBDocumentChunk.document_id == document.id
        ).all()

        if not chunks:
            return

        try:
            # Batch generate embeddings
            texts = [chunk.content for chunk in chunks]
            embeddings = await ai_service.create_embeddings_batch(texts)

            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
                chunk.embedding_model = settings.OPENAI_EMBEDDING_MODEL

            db.commit()
        except Exception as e:
            logger.error(f"Error generating chunk embeddings: {e}")

    async def _update_document_embedding(
        self, db: Session, document: KBDocument
    ) -> None:
        """Generate document-level embedding"""
        try:
            text = f"{document.title}\n"
            if document.description:
                text += f"{document.description}\n"
            if document.summary:
                text += document.summary

            embedding = await ai_service.create_embedding(text)
            document.embedding = embedding
            document.embedding_model = settings.OPENAI_EMBEDDING_MODEL
            db.commit()
        except Exception as e:
            logger.error(f"Error generating document embedding: {e}")

    async def get_documents_paginated(
        self,
        db: Session,
        page: int = 1,
        limit: int = 20,
        category_id: Optional[int] = None,
        search: Optional[str] = None,
        is_active: Optional[bool] = None,
        is_processed: Optional[bool] = None,
        lounge_id: Optional[int] = None,
        include_global: bool = True
    ) -> Tuple[List[KBDocument], int]:
        """Get paginated documents with filters"""
        query = db.query(KBDocument)

        if category_id:
            query = query.filter(KBDocument.category_id == category_id)
        if is_active is not None:
            query = query.filter(KBDocument.is_active == is_active)
        if is_processed is not None:
            query = query.filter(KBDocument.is_processed == is_processed)
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                or_(
                    KBDocument.title.ilike(search_term),
                    KBDocument.original_filename.ilike(search_term),
                    KBDocument.description.ilike(search_term)
                )
            )

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBDocument.lounge_id == lounge_id, KBDocument.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBDocument.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBDocument.lounge_id.is_(None))

        total = query.count()
        skip = (page - 1) * limit
        documents = query.order_by(KBDocument.created_at.desc()).offset(skip).limit(limit).all()

        return documents, total

    async def get_document(self, db: Session, document_id: int) -> Optional[KBDocument]:
        """Get a document by ID"""
        return db.query(KBDocument).filter(KBDocument.id == document_id).first()

    async def update_document(
        self, db: Session, document_id: int, data: dict
    ) -> Optional[KBDocument]:
        """Update a document's metadata"""
        document = await self.get_document(db, document_id)
        if not document:
            return None

        for key, value in data.items():
            if value is not None:
                setattr(document, key, value)

        db.commit()
        db.refresh(document)
        return document

    async def delete_document(self, db: Session, document_id: int) -> bool:
        """Delete a document and its file"""
        document = await self.get_document(db, document_id)
        if not document:
            return False

        # Store file info before deletion (needed for physical file cleanup)
        file_id = document.file_id
        file_record = db.query(File).filter(File.id == file_id).first()
        storage_path = file_record.storage_path if file_record else None

        # Delete the document first (cascades to delete chunks)
        db.delete(document)
        db.commit()

        # Now delete the File record from database
        if file_record:
            db.delete(file_record)
            db.commit()

        # Finally delete the physical file from storage
        if storage_path:
            try:
                await file_service.delete_file_by_path(storage_path)
            except Exception as e:
                logger.error(f"Error deleting physical file: {e}")

        return True

    async def reprocess_document(
        self, db: Session, document_id: int
    ) -> Optional[KBDocument]:
        """Reprocess a document"""
        document = await self.get_document(db, document_id)
        if not document:
            return None

        # Clear existing chunks
        db.query(KBDocumentChunk).filter(
            KBDocumentChunk.document_id == document_id
        ).delete()
        db.commit()

        # Reprocess
        await self._process_document(db, document)
        db.refresh(document)
        return document

    async def get_document_download_url(
        self, db: Session, document_id: int
    ) -> Optional[str]:
        """Get presigned download URL for document"""
        document = await self.get_document(db, document_id)
        if not document:
            return None

        return await file_service.get_file_url(document.file_id, db)

    # ============== FAQ Operations ==============

    async def create_faq(
        self, db: Session, data: dict, admin_id: int
    ) -> KBFaq:
        """Create a new FAQ and generate embedding"""
        faq = KBFaq(**data, created_by_id=admin_id)
        db.add(faq)
        db.commit()
        db.refresh(faq)

        if faq.is_included_in_rag:
            await self._update_faq_embedding(db, faq)

        return faq

    async def _update_faq_embedding(self, db: Session, faq: KBFaq) -> None:
        """Generate embedding for FAQ (question + answer)"""
        try:
            text = f"Question: {faq.question}\nAnswer: {faq.answer}"
            embedding = await ai_service.create_embedding(text)
            faq.embedding = embedding
            faq.embedding_model = settings.OPENAI_EMBEDDING_MODEL
            db.commit()
        except Exception as e:
            logger.error(f"Error generating FAQ embedding: {e}")

    async def get_faqs_paginated(
        self,
        db: Session,
        page: int = 1,
        limit: int = 20,
        category_id: Optional[int] = None,
        search: Optional[str] = None,
        is_active: Optional[bool] = None,
        lounge_id: Optional[int] = None,
        include_global: bool = True
    ) -> Tuple[List[KBFaq], int]:
        """Get paginated FAQs with filters"""
        query = db.query(KBFaq)

        if category_id:
            query = query.filter(KBFaq.category_id == category_id)
        if is_active is not None:
            query = query.filter(KBFaq.is_active == is_active)
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                or_(
                    KBFaq.question.ilike(search_term),
                    KBFaq.answer.ilike(search_term)
                )
            )

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBFaq.lounge_id == lounge_id, KBFaq.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBFaq.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBFaq.lounge_id.is_(None))

        total = query.count()
        skip = (page - 1) * limit
        faqs = query.order_by(KBFaq.sort_order, KBFaq.created_at.desc()).offset(skip).limit(limit).all()

        return faqs, total

    async def get_faq(self, db: Session, faq_id: int) -> Optional[KBFaq]:
        """Get a FAQ by ID"""
        return db.query(KBFaq).filter(KBFaq.id == faq_id).first()

    async def update_faq(
        self, db: Session, faq_id: int, data: dict
    ) -> Optional[KBFaq]:
        """Update a FAQ"""
        faq = await self.get_faq(db, faq_id)
        if not faq:
            return None

        content_changed = False
        for key, value in data.items():
            if value is not None:
                if key in ['question', 'answer']:
                    content_changed = True
                setattr(faq, key, value)

        db.commit()
        db.refresh(faq)

        if content_changed and faq.is_included_in_rag:
            await self._update_faq_embedding(db, faq)

        return faq

    async def delete_faq(self, db: Session, faq_id: int) -> bool:
        """Delete a FAQ"""
        faq = await self.get_faq(db, faq_id)
        if not faq:
            return False

        db.delete(faq)
        db.commit()
        return True

    # ============== RAG / Search Operations ==============

    async def semantic_search(
        self,
        db: Session,
        query: str,
        entity_types: Optional[List[str]] = None,
        category_ids: Optional[List[int]] = None,
        lounge_id: Optional[int] = None,
        include_global: bool = True,
        limit: int = 10
    ) -> List[Dict[str, Any]]:
        """Perform semantic search across KB entities"""
        # Generate query embedding
        query_embedding = await ai_service.create_embedding(query)

        results = []

        if not entity_types or "prompts" in entity_types:
            prompts = await self._search_prompts(
                db, query_embedding, category_ids, lounge_id, include_global, limit
            )
            results.extend(prompts)

        if not entity_types or "documents" in entity_types:
            docs = await self._search_documents(
                db, query_embedding, category_ids, lounge_id, include_global, limit
            )
            results.extend(docs)

        if not entity_types or "faqs" in entity_types:
            faqs = await self._search_faqs(
                db, query_embedding, category_ids, lounge_id, include_global, limit
            )
            results.extend(faqs)

        # Sort by similarity score and return top results
        results.sort(key=lambda x: x["similarity_score"], reverse=True)
        return results[:limit]

    async def _search_prompts(
        self, db: Session, query_embedding: List[float],
        category_ids: Optional[List[int]],
        lounge_id: Optional[int], include_global: bool, limit: int
    ) -> List[Dict[str, Any]]:
        """Search prompts by embedding similarity"""
        query = db.query(KBPrompt).filter(
            KBPrompt.is_active == True,
            KBPrompt.embedding.isnot(None)
        )
        if category_ids:
            query = query.filter(KBPrompt.category_id.in_(category_ids))

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBPrompt.lounge_id == lounge_id, KBPrompt.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBPrompt.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBPrompt.lounge_id.is_(None))

        prompts = query.all()

        # Calculate similarities
        candidates = [(p.id, p.embedding) for p in prompts if p.embedding]
        if not candidates:
            return []

        similar = ai_service.find_most_similar(query_embedding, candidates, limit)

        results = []
        for prompt_id, score in similar:
            prompt = next((p for p in prompts if p.id == prompt_id), None)
            if prompt:
                results.append({
                    "entity_type": "prompt",
                    "entity_id": prompt.id,
                    "title": prompt.title,
                    "content_preview": prompt.content[:200],
                    "similarity_score": score,
                    "category_name": prompt.category.name if prompt.category else None
                })

        return results

    async def _search_documents(
        self, db: Session, query_embedding: List[float],
        category_ids: Optional[List[int]],
        lounge_id: Optional[int], include_global: bool, limit: int
    ) -> List[Dict[str, Any]]:
        """Search documents by embedding similarity (using chunks)"""
        # Search through document chunks for more granular results
        query = db.query(KBDocumentChunk).join(KBDocument).filter(
            KBDocument.is_active == True,
            KBDocumentChunk.embedding.isnot(None)
        )
        if category_ids:
            query = query.filter(KBDocument.category_id.in_(category_ids))

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBDocument.lounge_id == lounge_id, KBDocument.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBDocument.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBDocument.lounge_id.is_(None))

        chunks = query.all()

        candidates = [(c.id, c.embedding) for c in chunks if c.embedding]
        if not candidates:
            return []

        similar = ai_service.find_most_similar(query_embedding, candidates, limit)

        results = []
        seen_docs = set()
        for chunk_id, score in similar:
            chunk = next((c for c in chunks if c.id == chunk_id), None)
            if chunk and chunk.document_id not in seen_docs:
                doc = chunk.document
                seen_docs.add(doc.id)
                results.append({
                    "entity_type": "document",
                    "entity_id": doc.id,
                    "title": doc.title,
                    "content_preview": chunk.content[:200],
                    "similarity_score": score,
                    "category_name": doc.category.name if doc.category else None
                })

        return results

    async def _search_faqs(
        self, db: Session, query_embedding: List[float],
        category_ids: Optional[List[int]],
        lounge_id: Optional[int], include_global: bool, limit: int
    ) -> List[Dict[str, Any]]:
        """Search FAQs by embedding similarity"""
        query = db.query(KBFaq).filter(
            KBFaq.is_active == True,
            KBFaq.embedding.isnot(None)
        )
        if category_ids:
            query = query.filter(KBFaq.category_id.in_(category_ids))

        # Lounge filtering
        if lounge_id is not None:
            if include_global:
                query = query.filter(
                    or_(KBFaq.lounge_id == lounge_id, KBFaq.lounge_id.is_(None))
                )
            else:
                query = query.filter(KBFaq.lounge_id == lounge_id)
        elif not include_global:
            query = query.filter(KBFaq.lounge_id.is_(None))

        faqs = query.all()

        candidates = [(f.id, f.embedding) for f in faqs if f.embedding]
        if not candidates:
            return []

        similar = ai_service.find_most_similar(query_embedding, candidates, limit)

        results = []
        for faq_id, score in similar:
            faq = next((f for f in faqs if f.id == faq_id), None)
            if faq:
                results.append({
                    "entity_type": "faq",
                    "entity_id": faq.id,
                    "title": faq.question[:100],
                    "content_preview": faq.answer[:200],
                    "similarity_score": score,
                    "category_name": faq.category.name if faq.category else None
                })

        return results

    async def get_rag_context(
        self,
        db: Session,
        query: str,
        max_items: int = 5,
        entity_types: Optional[List[str]] = None,
        lounge_id: Optional[int] = None,
        include_global: bool = True,
        similarity_threshold: float = 0.7
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """Get RAG context for AI chat"""
        results = await self.semantic_search(
            db=db,
            query=query,
            entity_types=entity_types,
            lounge_id=lounge_id,
            include_global=include_global,
            limit=max_items
        )

        # Build context string
        context_parts = []
        sources = []

        for item in results:
            if item["similarity_score"] >= similarity_threshold:
                context_parts.append(
                    f"[{item['entity_type'].upper()}] {item['title']}:\n{item['content_preview']}"
                )
                sources.append({
                    "type": item["entity_type"],
                    "id": item["entity_id"],
                    "title": item["title"]
                })

        context = "\n\n---\n\n".join(context_parts)
        return context, sources

    # ============== Stats Operations ==============

    async def get_stats(
        self, db: Session, lounge_id: Optional[int] = None, include_global: bool = True
    ) -> Dict[str, int]:
        """Get KB statistics, optionally filtered by lounge"""
        # Build lounge filter helper
        def lounge_filter(model, query):
            if lounge_id is not None:
                if include_global:
                    return query.filter(
                        or_(model.lounge_id == lounge_id, model.lounge_id.is_(None))
                    )
                return query.filter(model.lounge_id == lounge_id)
            elif not include_global:
                return query.filter(model.lounge_id.is_(None))
            return query

        cat_query = lounge_filter(KBCategory, db.query(func.count(KBCategory.id)))
        prompt_query = lounge_filter(KBPrompt, db.query(func.count(KBPrompt.id)))
        doc_query = lounge_filter(KBDocument, db.query(func.count(KBDocument.id)))
        faq_query = lounge_filter(KBFaq, db.query(func.count(KBFaq.id)))

        return {
            "total_categories": cat_query.scalar() or 0,
            "total_prompts": prompt_query.scalar() or 0,
            "total_documents": doc_query.scalar() or 0,
            "total_faqs": faq_query.scalar() or 0,
            "prompts_with_embeddings": lounge_filter(
                KBPrompt, db.query(func.count(KBPrompt.id)).filter(KBPrompt.embedding.isnot(None))
            ).scalar() or 0,
            "documents_with_embeddings": lounge_filter(
                KBDocument, db.query(func.count(KBDocument.id)).filter(KBDocument.embedding.isnot(None))
            ).scalar() or 0,
            "faqs_with_embeddings": lounge_filter(
                KBFaq, db.query(func.count(KBFaq.id)).filter(KBFaq.embedding.isnot(None))
            ).scalar() or 0,
            "total_document_chunks": db.query(func.count(KBDocumentChunk.id)).scalar() or 0,
            "unprocessed_documents": lounge_filter(
                KBDocument, db.query(func.count(KBDocument.id)).filter(KBDocument.is_processed == False)
            ).scalar() or 0
        }

    # ============== Bulk Operations ==============

    async def regenerate_all_embeddings(
        self, db: Session, entity_type: Optional[str] = None
    ) -> Dict[str, int]:
        """Regenerate embeddings for all KB content"""
        results = {"prompts": 0, "documents": 0, "faqs": 0}

        if not entity_type or entity_type == "prompts":
            prompts = db.query(KBPrompt).filter(
                KBPrompt.is_included_in_rag == True
            ).all()
            for prompt in prompts:
                await self._update_prompt_embedding(db, prompt)
                results["prompts"] += 1

        if not entity_type or entity_type == "documents":
            documents = db.query(KBDocument).filter(
                KBDocument.is_processed == True
            ).all()
            for doc in documents:
                await self._generate_chunk_embeddings(db, doc)
                await self._update_document_embedding(db, doc)
                results["documents"] += 1

        if not entity_type or entity_type == "faqs":
            faqs = db.query(KBFaq).filter(
                KBFaq.is_included_in_rag == True
            ).all()
            for faq in faqs:
                await self._update_faq_embedding(db, faq)
                results["faqs"] += 1

        return results


# Singleton instance
knowledge_base_service = KnowledgeBaseService()
