#!/usr/bin/env python3
"""Generate Prompterly Security & Data Architecture Standard v2 as DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_section_header(doc, num, title):
    """Add a peach-coloured section header (matching original PDF style)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    shade_cell(cell, "FEE2C5")  # Peach background

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {num}. {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


def add_subsection(doc, num, title):
    """Add a grey subsection header."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F0F0F0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {num} {title}")
    run.bold = True
    run.font.size = Pt(11)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    if not p.runs:
        run = p.add_run(text)
    else:
        # The bullet style doesn't add a run automatically; ensure text is set
        p.add_run(text)
    for r in p.runs:
        r.font.size = Pt(10)


def add_bullet_simple(doc, text):
    """Bullet using a dash for consistency with original PDF."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"• {text}")
    run.font.size = Pt(10)


def add_code_box(doc, text):
    """Add a monospace code/example block."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}")
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)
    doc.add_paragraph()


def build():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== Cover Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Prompterly Security & Data\nArchitecture Standard")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(26, 26, 31)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ["Version 2.0", "Last Updated: April 2026", "Next Review: October 2026"]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    add_para(doc,
        "This document outlines the minimum security, privacy, backup, and data protection "
        "requirements for the Prompterly platform before public launch. The goal is to protect "
        "user data, mentor intellectual property, and platform operations while complying with "
        "applicable privacy laws and industry best practices. This document applies to:")

    for item in [
        "Platform infrastructure",
        "User data",
        "Mentor content",
        "AI coaching lounge interactions",
        "Payment processing through Stripe",
        "Third-party data processors and sub-processors",
    ]:
        add_bullet_simple(doc, item)

    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(
        "Prompterly is a platform that will store deeply personal reflections. It is "
        "expected that users will treat the coaching lounge like a journal. That creates a "
        "psychological trust obligation, even if the law doesn't strictly require it. "
        "Therefore, the goal should be security posture closer to a journaling app than a "
        "simple chatbot.")
    r2.font.size = Pt(10)

    add_heading(doc, "Document Control:")
    for item in [
        "Owner: Prompterly Compliance / Engineering Team",
        "Review Cadence: Annual or upon major regulatory change",
        "Distribution: Internal (engineering, mentors, admin staff)",
    ]:
        add_bullet_simple(doc, item)

    doc.add_page_break()

    # ===== Section 1 =====
    add_section_header(doc, "1", "Data Categories Stored on Prompterly")
    add_para(doc, "Prompterly stores several categories of information which must be protected appropriately.")

    add_heading(doc, "User data")
    for item in ["Name", "Email address", "Account settings", "Notification preferences", "Language and timezone preferences"]:
        add_bullet_simple(doc, item)

    add_heading(doc, "User generated content")
    for item in ["Coaching lounge chat messages", "Notebook entries", "Time capsule entries"]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Platform data")
    for item in ["Subscription records", "Usage logs", "Session tokens", "System diagnostics", "Audit logs"]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Mentor data")
    for item in ["Onboarding framework materials and IP", "Coaching lounge configuration", "Mentor profile information", "Mentor framework version history"]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Payment data")
    for item in ["Stripe subscription identifiers", "Payment status records"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Prompterly does not store full credit card details, which are handled entirely by Stripe.")

    # ===== Section 2 =====
    add_section_header(doc, "2", "Data Separation & Pseudonymisation Architecture")
    add_para(doc, "Prompterly must store user identity information separately from user-generated content.")
    add_para(doc, "This architecture reduces the risk of unnecessary exposure of sensitive conversations and aligns with privacy-by-design principles.")

    add_heading(doc, "Core Principle")
    add_para(doc, "User conversations and personal reflections should not be directly linked to identifiable personal information in the same database structures. Instead, content must be associated with an internal pseudonymous identifier.")

    add_heading(doc, "Implementation Note")
    add_para(doc, "This architecture does not require multiple physical databases. Logical separation within the schema is acceptable provided:")
    for item in [
        "Identity fields are stored in dedicated tables",
        "Conversation and content tables reference user_uuid only",
        "Personal identifiers are not duplicated across content tables",
    ]:
        add_bullet_simple(doc, item)

    add_subsection(doc, "2.1", "Internal User Identifier")
    add_para(doc, "Each user account must be assigned a unique internal identifier.")
    add_code_box(doc, "user_uuid")
    add_para(doc, "This identifier will be used as the primary reference key across platform systems. The internal identifier must not contain personal information.")

    add_subsection(doc, "2.2", "Separation of Identity Data and Content Data")
    add_para(doc, "Prompterly must store data across separate logical domains.")

    add_heading(doc, "Identity Data (Personally Identifiable Information)")
    for item in ["Name", "Email address", "Authentication credentials", "Stripe customer ID", "Account preferences", "Notification settings"]:
        add_bullet_simple(doc, item)
    add_para(doc, "This data should be stored in the User Identity table or equivalent service.")

    add_heading(doc, "User Generated Content")
    add_para(doc, "The following data must reference user_uuid only, not personal identifiers.")
    for item in ["Chat messages", "Chat threads", "Notebook entries", "Time capsule entries", "AI coaching lounge interaction records"]:
        add_bullet_simple(doc, item)
    add_para(doc, "These tables must not contain name, email, or other identifying fields.")

    add_subsection(doc, "2.3", "Mapping Between Identity and Content")
    add_para(doc, "A mapping relationship must exist between:")
    add_code_box(doc, "user_uuid -> identity record")
    add_para(doc, "This mapping must be stored in the identity layer and should not be duplicated across other tables.")
    add_para(doc, "Access to this mapping should be restricted to only the services that require it for legitimate platform functions.")

    add_subsection(doc, "2.4", "Access Restrictions")
    add_para(doc, "Not all services or administrators should be able to directly associate user conversations with identity information. Administrative tools capable of linking identity data with conversation data must be restricted to authorised personnel only and protected with multi-factor authentication.")
    add_para(doc, "Best practice is that:")
    for item in [
        "Standard application services operate using user_uuid",
        "Identity lookup is performed only when necessary (e.g. login, account management)",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Where possible:")
    for item in [
        "Developer debugging tools should operate on pseudonymous identifiers",
        "Logs should avoid storing personal identifiers alongside conversation data",
    ]:
        add_bullet_simple(doc, item)

    add_subsection(doc, "2.5", "Logging and Monitoring Considerations")
    add_para(doc, "System logs must avoid storing sensitive conversation content or unnecessary personal identifiers.")
    add_para(doc, "Logs may include:")
    for item in ["user_uuid", "timestamps", "request metadata"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Logs should not include full conversation text unless required for debugging or safety monitoring.")

    add_subsection(doc, "2.6", "Data Export and User Access")
    add_para(doc, "When users request to export their data, the system may temporarily reconstruct the relationship between:")
    for item in ["identity data", "conversation data"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Exports must include:")
    for item in ["chat history", "notebook entries", "time capsule entries"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Export formats must include JSON, CSV, and PDF.")

    add_subsection(doc, "2.7", "Data Deletion")
    add_para(doc, "When a user deletes a conversation thread, notebook entry, or time capsule entry:")
    add_bullet_simple(doc, "The content should be removed from the content tables associated with their user_uuid.")
    add_para(doc, "When a user deletes their account:")
    for item in [
        "Personal identity information should be anonymised",
        "Authentication credentials revoked",
        "Content records remain pseudonymous unless administratively deleted",
    ]:
        add_bullet_simple(doc, item)

    add_subsection(doc, "2.8", "Benefits of This Architecture")
    add_para(doc, "This design provides several important protections:")
    for item in [
        "Reduces unnecessary exposure of personal identity in content systems",
        "Limits the number of systems that can link identity to conversations",
        "Supports privacy-by-design principles",
        "Reduces the impact of accidental data exposure",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "This architecture does not eliminate legal disclosure obligations, but it ensures that access to identifiable user data is appropriately controlled.")

    # ===== Section 3 =====
    add_section_header(doc, "3", "Encryption Standards")
    add_para(doc, "All user data must be protected using encryption both when stored and when transmitted. All production databases must not be publicly accessible on the internet and must be restricted to application servers via private network rules. Production databases must only accept connections from authorised application servers through private network rules or security groups.")

    add_heading(doc, "Encryption In Transit")
    add_para(doc, "All traffic between users and the Prompterly platform must be encrypted using HTTPS with TLS 1.2 or higher. HTTP connections must automatically redirect to HTTPS. TLS certificates must be valid and configured to prevent insecure fallback connections.")

    add_heading(doc, "Encryption At Rest")
    add_para(doc, "The following systems must use encryption at rest:")
    for item in ["Platform database (MySQL or equivalent)", "Object storage (S3 or equivalent)", "Backup storage systems"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Industry standard AES-256 encryption (or equivalent cloud-provider managed encryption) should be used where available.")
    add_para(doc, "The mapping between user identity and user_uuid must be protected using the same encryption and access controls applied to other sensitive personal data.")

    add_heading(doc, "Sensitive Content Encryption")
    add_para(doc, "Prompterly must apply application-level encryption to highly sensitive user-generated content. The following fields must be encrypted before storage:")
    for item in ["Chat message content", "Notebook entry text", "Time capsule entry content"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Encryption must occur before the data is written to the database.")
    add_para(doc, "Encryption keys must be managed through a secure key management service (e.g., AWS KMS or equivalent) and must not be stored directly in the application database.")
    add_para(doc, "Decryption should occur only within the application layer when rendering user content to the authenticated user.")

    add_heading(doc, "Search requirements")
    add_para(doc, "Prompterly must preserve the user's ability to search within the current chat thread.")
    add_para(doc, "If chat message content is encrypted at the application layer, developers must implement a search method that allows thread-level search without exposing full plaintext content unnecessarily.")
    add_para(doc, "This may be achieved by:")
    for item in [
        "Decrypting the current thread in the application layer for search, or",
        "Maintaining a separate search index that does not store full message content in plain text",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 4 =====
    add_section_header(doc, "4", "Password & Authentication Security")
    add_para(doc, "Passwords must never be stored in plain text. Passwords must be stored using strong hashing methods such as bcrypt.")
    add_para(doc, "Authentication tokens must:")
    for item in [
        "Have a defined expiration time appropriate to the application session (e.g., hours rather than months)",
        "Be securely stored",
        "Be transmitted only over encrypted connections",
        "Be invalidated upon user logout where possible",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Administrative access must require:")
    for item in [
        "Role-based permissions",
        "Multi-factor authentication (MANDATORY for all admin accounts)",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Login & API Rate Limiting")
    add_para(doc, "Authentication endpoints and sensitive API routes must implement rate limiting to prevent automated abuse.")
    add_para(doc, "Minimum expectations:")
    for item in [
        "Limit repeated login attempts from the same IP address",
        "Temporarily block excessive failed login attempts",
        "Apply rate limiting to authentication and account recovery endpoints",
        "Per-account lockout after 5 failed attempts (15-minute cooldown)",
        "Notify the user via email when their account is locked",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Two-Factor Authentication (2FA) — NEW")
    add_para(doc, "Prompterly must offer 2FA for all users:")
    for item in [
        "Email OTP (default method — sends 6-digit code to user's email)",
        "Authenticator app (TOTP) as an optional alternative for advanced users",
        "MANDATORY for all admin and mentor accounts",
        "Cannot be disabled by the user once enforced for admin/mentor roles",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Account Recovery — NEW")
    add_para(doc, "If a user loses access to their MFA device or cannot receive emails:")
    for item in [
        "Account recovery must be available via verified support process",
        "Identity verification required before recovery (e.g., security questions, ID verification)",
        "Recovery actions must be logged in the audit trail",
        "User must be notified via all known contact methods when recovery is initiated",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 5 =====
    add_section_header(doc, "5", "Secrets & Credential Management")
    add_para(doc, "Sensitive credentials must not be hardcoded in application code.")
    add_para(doc, "This includes:")
    for item in ["Database credentials", "API keys", "Stripe secret keys", "Email service credentials", "AI provider keys", "Encryption keys"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Secrets must be stored using a secure secret management system such as:")
    for item in ["AWS Secrets Manager", "AWS Parameter Store", "Hashicorp Vault", "Or an equivalent managed secret storage"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Access to secrets must be restricted to only the services that require them.")

    # ===== Section 6 =====
    add_section_header(doc, "6", "Data Backups & Disaster Recovery")
    add_para(doc, "Prompterly must maintain reliable backup and recovery procedures.")

    add_heading(doc, "Database Backups")
    add_para(doc, "Automated backups of the primary database must occur at least once per day. Backups must be:")
    for item in [
        "Stored separately from the primary server",
        "Encrypted at rest using the same encryption standard applied to the primary database (AES-256 or equivalent)",
        "Stored in secure cloud storage (e.g., encrypted S3 bucket or equivalent)",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "If backups are stored in cloud object storage, server-side encryption must be enabled.")

    add_heading(doc, "Backup Retention")
    add_para(doc, "Recommended retention periods:")
    for item in ["Daily backups: 30 days", "Monthly backups: 12 months"]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Restore Capability")
    add_para(doc, "The system must support restoration of database backups in the event of:")
    for item in ["Infrastructure failure", "Accidental deletion", "Security incident"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Bitsclan / developers must test restoration procedures periodically (at least quarterly) to confirm backups are functional.")

    add_heading(doc, "Recovery Objectives — NEW")
    for item in [
        "Recovery Time Objective (RTO): 4 hours maximum downtime",
        "Recovery Point Objective (RPO): 24 hours maximum data loss (matches daily backup cadence)",
        "Service uptime target: 99.5%",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 7 =====
    add_section_header(doc, "7", "Data Retention Policy")
    add_para(doc, "Prompterly defines the following timeframes for how long different categories of data are stored:")
    for item in [
        "Chat Messages — Indefinite: Retained until the user deletes the associated conversation thread",
        "Notebook Entries — Indefinite: Retained until the user deletes the entry",
        "Time Capsule Entries — Indefinite: Retained until the user deletes the entry",
        "System Logs — 60 days, then automatically rotated and deleted unless required for investigation or legal preservation",
        "Audit Logs — At least 12 months",
        "Backups — Daily backups: 30 days, monthly backups: 12 months",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Log Rotation")
    add_para(doc, "Operational logs must implement automatic log rotation. Minimum requirements:")
    for item in [
        "Logs must rotate daily or by size threshold",
        "Rotated logs must be retained for 60 days",
        "Older logs must be automatically deleted unless preserved under legal hold",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Legal Retention Override — NEW")
    add_para(doc, "Where Australian or applicable law requires longer retention (e.g. tax records: 5 years per ATO requirements; payment records: 7 years), legal retention requirements override the policy above. The longer of the two periods applies.")

    add_heading(doc, "Account Deletion")
    add_para(doc, "When a user deletes their account:")
    for item in [
        "Direct identifiers (name, email, Stripe ID) must be removed or irreversibly anonymised",
        "Authentication credentials must be revoked",
        "Content records remain associated only with the internal user_uuid unless administratively deleted",
        "A confirmation email must be sent to the user before deletion is finalised",
        "A 30-day cooling off period applies (user can cancel deletion within this window)",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 8 =====
    add_section_header(doc, "8", "Legal Preservation (Legal Hold)")
    add_para(doc, "Prompterly must maintain the ability to temporarily preserve account data when required.")
    add_para(doc, "If a legal request or investigation occurs, the platform must be able to:")
    for item in [
        "Temporarily disable deletion of account data",
        "Preserve relevant records for legal review",
        "Document the reason and authorising party for the hold",
        "Notify affected users where legally permitted",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "This capability should be controlled through an internal administrative mechanism.")

    # ===== Section 9 =====
    add_section_header(doc, "9", "Audit Logging")
    add_para(doc, "Prompterly must maintain logs of critical platform actions.")
    add_para(doc, "The system should record:")
    for item in [
        "Account creation and deletion",
        "Administrative actions",
        "Data deletion requests",
        "Subscription changes",
        "Security-related events",
        "Consent changes (opt-in/opt-out)",
        "Permission and role changes",
        "Data exports",
        "Legal hold actions",
        "Mentor framework version changes",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Audit logs must be stored separately from standard application logs where possible.")
    add_para(doc, "Audit logs must be tamper-evident (e.g., append-only or cryptographically signed) where feasible.")

    # ===== Section 10 =====
    add_section_header(doc, "10", "User Data Export")
    add_para(doc, "Users must have the ability to export their data upon request.")
    add_para(doc, "Exportable data may include:")
    for item in ["Chat conversation history", "Notebook entries", "Time capsule entries", "Profile information", "Subscription and billing history"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Exports must be provided in formats such as:")
    for item in ["JSON", "CSV", "PDF (human-readable)"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Export features should avoid including mentor branding where the lounge has been removed due to safety or legal concerns.")

    add_heading(doc, "Response Timeframes — NEW")
    for item in [
        "GDPR Subject Access Requests: 30 days from verified request",
        "CCPA Consumer Requests: 45 days (extendable by 45 days with notice)",
        "Australian Privacy Act access requests: 30 days",
        "Identity verification required before fulfilling requests",
        "First request per user is free of charge",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 11 =====
    add_section_header(doc, "11", "Payment Processing (Stripe)")
    add_para(doc, "Prompterly uses Stripe for payment processing.")

    add_heading(doc, "Stripe Responsibilities")
    add_para(doc, "Stripe securely handles:")
    for item in ["Credit card storage", "Payment authorisation", "Payment processing"]:
        add_bullet_simple(doc, item)
    add_para(doc, "Prompterly must never store raw credit card numbers.")

    add_heading(doc, "Prompterly Responsibilities")
    add_para(doc, "Prompterly stores limited Stripe-related records such as:")
    for item in ["Stripe customer ID", "Subscription ID", "Payment status"]:
        add_bullet_simple(doc, item)
    add_para(doc, "These identifiers must be stored securely within the platform database.")

    add_heading(doc, "Stripe Webhooks")
    add_para(doc, "Stripe webhook endpoints must:")
    for item in [
        "Validate webhook signatures",
        "Reject requests without valid signatures",
        "Enforce timestamp tolerance to prevent replay attacks",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 12 =====
    add_section_header(doc, "12", "Minimum Age Requirements")
    add_para(doc, "Prompterly is intended for users aged 18 years or older. Account creation must require users to confirm:")
    add_code_box(doc, '"I confirm that I am at least 18 years of age or older"')

    add_heading(doc, "Underage Account Discovery — NEW")
    add_para(doc, "If Prompterly discovers a user is under 18:")
    for item in [
        "The account must be immediately suspended",
        "All personal data must be deleted within 30 days",
        "Parents/guardians may be notified where appropriate and lawful",
        "A reporting mechanism must exist for users to flag suspected underage accounts",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 13 =====
    add_section_header(doc, "13", "Security Monitoring & Incident Response")
    add_para(doc, "Prompterly must maintain procedures for responding to security incidents. If a data breach or security incident occurs, Prompterly must:")
    for item in [
        "Investigate the issue promptly",
        "Secure affected systems",
        "Notify affected users where legally required",
        "Document the incident in the breach log",
        "Conduct a root cause analysis",
        "Apply remediation and preventive measures",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Notifiable Data Breaches (NDB) Scheme — NEW")
    add_para(doc, "Prompterly is subject to Australia's Notifiable Data Breaches (NDB) scheme under the Privacy Act 1988. If a data breach is likely to result in serious harm to affected individuals, Prompterly must:")
    for item in [
        "Assess the breach within 30 days of becoming aware",
        "Notify the Office of the Australian Information Commissioner (OAIC) as soon as practicable",
        "Notify affected individuals as soon as practicable",
        "Provide the OAIC and users with details: what happened, what data was affected, recommended actions",
        "Maintain a breach log with date, scope, severity, response actions, and outcomes",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Breach Notification Content")
    add_para(doc, "Notifications to affected users must include:")
    for item in [
        "Description of what happened (in plain language)",
        "Type of personal information involved",
        "Recommended steps the user should take (e.g., change password)",
        "Contact information for support",
        "Reference number for the incident",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Detection & Monitoring")
    for item in [
        "Anomaly detection on authentication and admin actions",
        "Failed login monitoring with alerting",
        "Regular review of audit logs",
        "Vulnerability scanning of dependencies (npm, pip)",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 14 =====
    add_section_header(doc, "14", "AI System Data Boundaries")
    add_para(doc, "Prompterly's AI coaching lounges generate responses using:")
    for item in ["Mentor-provided frameworks", "User prompts", "System guardrails"]:
        add_bullet_simple(doc, item)
    add_para(doc, "User conversations must not be used to train public AI models.")
    add_para(doc, "AI responses are generated in real time and are not intended to replace professional services such as therapy, legal advice, or medical advice.")

    add_heading(doc, "AI Disclosure Requirements — NEW")
    add_para(doc, "Under the Australian Consumer Law and AI Ethics Framework:")
    for item in [
        "Users must be clearly informed they are interacting with an AI, not a human",
        "This disclosure must appear on first interaction with each coaching lounge",
        "AI personas must not be presented in a way that misleads users into believing they are human",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "AI Safety Disclaimers")
    add_para(doc, "The following disclaimer must be visible to users:")
    add_code_box(doc, '"This is an AI coaching assistant, not a licensed professional."')
    add_code_box(doc, '"Do not use this service for medical, legal, or financial advice."')
    add_code_box(doc, '"In a crisis, contact emergency services or a qualified practitioner."')

    add_heading(doc, "Crisis Intervention Protocol — NEW")
    add_para(doc, "Because Prompterly stores deeply personal reflections, the system must implement basic safety mechanisms for users in crisis:")
    for item in [
        "Detect keywords or phrases related to self-harm, suicide, abuse, or imminent danger",
        "Display a crisis support message overriding the normal AI response",
        "Provide region-appropriate emergency contact information",
        "Australia: Lifeline 13 11 14, Beyond Blue 1300 22 4636, 000 for emergencies",
        "UK: Samaritans 116 123, 999 for emergencies",
        "USA: 988 Suicide & Crisis Lifeline, 911 for emergencies",
        "Log the event (without exposing the user's content) for safety review",
        "Never provide medical, psychological, or therapeutic advice in response to crisis content",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 15 =====
    add_section_header(doc, "15", "Mentor Framework / IP Versioning")
    add_para(doc, "Prompterly must maintain versioned configuration records for each Coaching Lounge framework and AI system instruction set.")
    add_para(doc, "Configuration includes, but is not limited to:")
    for item in [
        "Mentor frameworks",
        "Prompt templates",
        "AI system instructions",
        "Behavioural guardrails",
        "Tone or persona configuration",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "When a mentor or administrator modifies any of the above, the system must:")
    for item in [
        "Create a new configuration version",
        "Preserve all previous versions unchanged (versions must be immutable)",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Existing configuration versions must not be overwritten.")
    add_para(doc, "Each chat session must record the configuration version used to generate responses so that past interactions can be reconstructed if required.")
    add_para(doc, "The platform must support rollback to a previous configuration version if a configuration update produces unintended behaviour. Rollback must not modify or delete historical versions.")

    add_heading(doc, "Mentor IP Ownership — NEW")
    add_para(doc, "Prompterly must clearly define ownership of mentor-provided frameworks:")
    for item in [
        "Mentors retain ownership of their original framework IP",
        "Prompterly receives a licence to use the framework for the duration of the mentor's agreement",
        "If a mentor leaves the platform, their frameworks must be archived (not deleted) for a defined period",
        "Mentors must be able to export their own framework configurations on request",
        "Confidentiality of mentor-provided training data must be maintained",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 16 =====
    add_section_header(doc, "16", "Data Protection Compliance")
    add_para(doc, "Prompterly will operate globally and must comply with relevant privacy frameworks including:")
    for item in [
        "Australian Privacy Act 1988 (APPs)",
        "Australian Notifiable Data Breaches (NDB) Scheme",
        "GDPR (European Union)",
        "CCPA/CPRA (California)",
        "UK GDPR",
        "New Zealand Privacy Act 2020",
        "Canada PIPEDA",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Prompterly's privacy policy must clearly disclose:")
    for item in [
        "What data is collected",
        "How it is used",
        "How long it is retained",
        "How users may request deletion or export",
        "Who data is shared with (third parties)",
        "Cross-border data transfers",
        "User rights and how to exercise them",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Privacy Policy Versioning — NEW")
    for item in [
        "The privacy policy must include a 'last updated' date",
        "Policy versions must be tracked",
        "Users must be re-prompted to accept the policy when material changes occur",
        "Records of which version each user accepted must be retained",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 17 (NEW) =====
    add_section_header(doc, "17", "Consent Management (NEW)")
    add_para(doc, "Prompterly must implement granular and revocable consent mechanisms.")

    add_heading(doc, "Consent Collection")
    for item in [
        "Explicit opt-in at signup for privacy policy and terms of service",
        "Separate consent for marketing communications (not pre-checked)",
        "Separate consent for analytics and tracking (where applicable)",
        "Age confirmation (18+) explicit consent required",
        "All consent must be timestamped and recorded in the audit trail",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Consent Withdrawal")
    for item in [
        "Users must be able to withdraw consent at any time from account settings",
        "Withdrawing consent must not require contacting support",
        "Withdrawal must take effect within 24 hours",
        "Users must be informed of consequences (e.g., account closure if essential consent withdrawn)",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Re-consent")
    add_para(doc, "When privacy policy or terms of service are materially updated:")
    for item in [
        "Users must be prompted to accept the new version on next login",
        "Continued use without acceptance is not permitted for material changes",
        "The version number and acceptance timestamp must be recorded",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 18 (NEW) =====
    add_section_header(doc, "18", "Third-Party Processors & Cross-Border Transfers (NEW)")
    add_para(doc, "Prompterly uses third-party services that process user data. Under Australian Privacy Principle 8 and GDPR Article 28, these relationships must be disclosed and governed by formal agreements.")

    add_heading(doc, "Current Sub-Processors")
    for item in [
        "Anthropic (Claude AI) — USA — processes chat content for AI responses",
        "Amazon Web Services — hosts database, file storage, backups",
        "Stripe — USA — payment processing",
        "Postmark — USA — transactional email delivery",
        "Sentry (if used) — error monitoring",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Vendor Requirements")
    add_para(doc, "All sub-processors must:")
    for item in [
        "Have a signed Data Processing Agreement (DPA)",
        "Demonstrate equivalent or stronger privacy protections",
        "Be listed in Prompterly's privacy policy",
        "Provide breach notification within their own SLAs",
        "Support sub-processor change notifications to Prompterly",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Cross-Border Disclosure")
    add_para(doc, "Prompterly must disclose:")
    for item in [
        "Which countries user data is transferred to",
        "What categories of data are transferred",
        "That overseas recipients may not be subject to Australian privacy law",
        "Any safeguards in place (Standard Contractual Clauses, adequacy decisions, etc.)",
    ]:
        add_bullet_simple(doc, item)
    add_para(doc, "Users must consent to cross-border transfers as part of accepting the privacy policy.")

    # ===== Section 19 (NEW) =====
    add_section_header(doc, "19", "Privacy by Design (NEW)")
    add_para(doc, "Prompterly adopts a Privacy by Design approach as required under GDPR Article 25 and recommended by the OAIC.")

    add_heading(doc, "Core Principles")
    for item in [
        "Proactive, not reactive — anticipate privacy risks before they occur",
        "Privacy as the default setting — users should not have to opt in to protection",
        "Privacy embedded into design — not bolted on after the fact",
        "Full functionality — privacy and usability are not mutually exclusive",
        "End-to-end security — data protected throughout its lifecycle",
        "Visibility and transparency — users can verify privacy practices",
        "Respect for user privacy — keep user interests paramount",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Implementation Requirements")
    for item in [
        "Privacy review required for all new features",
        "Data Protection Impact Assessment (DPIA) for high-risk processing",
        "Default settings must minimise data collection",
        "Privacy-friendly defaults (e.g., notification preferences off by default for marketing)",
        "Documentation of data flows for each feature",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 20 (NEW) =====
    add_section_header(doc, "20", "Cookie Policy (NEW)")
    add_para(doc, "Where Prompterly uses cookies, web storage, or similar tracking technologies:")

    add_heading(doc, "Cookie Categories")
    for item in [
        "Essential cookies — required for the platform to function (session, authentication)",
        "Analytics cookies — measure usage patterns",
        "Marketing cookies — track for advertising purposes",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Consent Requirements")
    for item in [
        "Cookie consent banner displayed on first visit (especially for EU users under GDPR)",
        "Granular controls — users can accept/reject categories independently",
        "Essential cookies do not require consent but must be disclosed",
        "Cookie preferences can be changed at any time from settings",
        "Cookie policy must list all cookies, their purpose, and duration",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 21 (NEW) =====
    add_section_header(doc, "21", "Operational Security (NEW)")
    add_para(doc, "Beyond technical controls, Prompterly must maintain operational security practices.")

    add_heading(doc, "Staff Access Controls")
    for item in [
        "Principle of least privilege — staff only have access to what they need",
        "All production access logged and audited",
        "Off-boarding process: revoke access immediately when staff leave",
        "Quarterly access review",
        "Background checks for staff with access to sensitive data",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Security Testing")
    for item in [
        "Annual penetration test by qualified third party",
        "Quarterly automated vulnerability scans",
        "Dependency vulnerability scanning (npm audit, pip-audit) on every build",
        "Security patching within 30 days of patch availability (critical patches within 7 days)",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Training & Awareness")
    for item in [
        "Annual privacy and security training for all staff",
        "Onboarding training for new staff includes data protection",
        "Customer support trained on privacy requests and breach handling",
        "Incident response drills conducted at least annually",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Documentation")
    for item in [
        "All security policies documented and accessible to staff",
        "Records of Processing Activities (ROPA) maintained",
        "Data flow diagrams for major features",
        "Incident response runbook",
        "Recovery procedures runbook",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 22 (NEW) =====
    add_section_header(doc, "22", "Service Availability & SLAs (NEW)")
    add_heading(doc, "Uptime Targets")
    for item in [
        "Service uptime target: 99.5% (excluding planned maintenance)",
        "Planned maintenance windows: outside peak hours, with 48-hour notice",
        "Status page available for users to check incidents",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Recovery Objectives")
    for item in [
        "Recovery Time Objective (RTO): 4 hours maximum downtime in disaster scenarios",
        "Recovery Point Objective (RPO): 24 hours maximum data loss",
        "Failover procedures documented and tested annually",
    ]:
        add_bullet_simple(doc, item)

    # ===== Section 23 (NEW) =====
    add_section_header(doc, "23", "Document Control & Review (NEW)")
    add_heading(doc, "Version History")
    for item in [
        "Version 1.0 — Initial standard",
        "Version 2.0 — April 2026 — Added sections on consent management, third-party processors, AI safety, crisis intervention, breach notification specifics, privacy by design, cookies, operational security, SLAs, and document control",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Review Cadence")
    for item in [
        "Annual review (next: October 2026)",
        "Ad-hoc review upon major regulatory changes (e.g., new privacy law)",
        "Ad-hoc review after any significant security incident",
        "Updates require sign-off from compliance and engineering leads",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Distribution")
    for item in [
        "Internal: all staff, contractors, mentors with platform access",
        "External: shared with auditors, legal, and key partners under NDA",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Compliance Owner")
    for item in [
        "Primary owner: Prompterly Compliance Officer (to be designated)",
        "Technical owner: Engineering Lead (Bitsclan)",
        "Business owner: Lauren Donnelly",
    ]:
        add_bullet_simple(doc, item)

    output = "/home/wajid-ali/Projects/Prompterly/RESTAPI/RESTAPI-Prompterly/Prompterly_Security_and_Data_Architecture_Standard_v2.docx"
    doc.save(output)
    print(f"Generated: {output}")


if __name__ == "__main__":
    build()
